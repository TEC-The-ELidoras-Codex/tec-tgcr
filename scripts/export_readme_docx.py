from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document

REPO_ROOT = Path(__file__).resolve().parents[1]
README_PATH = REPO_ROOT / "README.md"
DEFAULT_OUTPUT = REPO_ROOT / "docs" / "exports" / "README.docx"


def add_paragraph(doc: Document, text: str, style: str | None = None) -> None:
    paragraph = doc.add_paragraph(text)
    if style:
        paragraph.style = style


def render_markdown(doc: Document, markdown: str) -> None:
    in_code_block = False
    code_lines: list[str] = []

    for raw_line in markdown.splitlines():
        line = raw_line.rstrip()

        if line.startswith("```"):
            if in_code_block:
                doc.add_paragraph("\n".join(code_lines), style="Intense Quote")
                code_lines.clear()
            in_code_block = not in_code_block
            continue

        if in_code_block:
            code_lines.append(line)
            continue

        if not line.strip():
            doc.add_paragraph("")
            continue

        if line.startswith("#"):
            level = len(line) - len(line.lstrip("#"))
            text = line[level:].strip()
            style = f"Heading {min(level, 9)}"
            add_paragraph(doc, text, style)
            continue

        if line.startswith(('- ', '* ')):
            add_paragraph(doc, line[2:].strip(), style="List Bullet")
            continue

        if line[0].isdigit() and line.split(".", 1)[0].isdigit():
            add_paragraph(doc, line, style="List Number")
            continue

        add_paragraph(doc, line)

    if code_lines:
        doc.add_paragraph("\n".join(code_lines), style="Intense Quote")


def export_readme(output: Path) -> Path:
    document = Document()
    markdown = README_PATH.read_text(encoding="utf-8")
    render_markdown(document, markdown)

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    return output


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Export README.md to DOCX")
    parser.add_argument(
        "--output",
        type=Path,
        default=DEFAULT_OUTPUT,
        help=f"Destination .docx file (default: {DEFAULT_OUTPUT.relative_to(REPO_ROOT)})",
    )
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    output_path = export_readme(args.output)
    rel = output_path.relative_to(REPO_ROOT)
    print(f"README exported to {rel}")


if __name__ == "__main__":
    main()
