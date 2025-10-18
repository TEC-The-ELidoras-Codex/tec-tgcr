#!/usr/bin/env python3
"""
Export comprehensive documentation compendium to DOCX format.

Usage:
    python scripts/export_compendium.py
    python scripts/export_compendium.py --output custom_path.docx
"""

import argparse
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE

# Manifest of documents to include in compendium
# (Title, Path) tuples
MANIFEST = [
    ("README", Path("README.md")),
    ("Agent Overview", Path("docs/AGENT_OVERVIEW.md")),
    ("Agent Persona: Airth Research Guard", Path("docs/AGENT_AIRTH.md")),
    ("Design: Spin Model", Path("docs/DESIGN_SPIN_MODEL.md")),
    ("WordPress Operations Guide", Path("docs/WORDPRESS_WPCOM_OPS.md")),
    ("GitHub Secrets Setup", Path("docs/GITHUB_SECRETS_SETUP.md")),
    ("Quick Secrets Fill Reference", Path("docs/QUICK_SECRETS_FILL.md")),
    ("Microsoft 365 Integration", Path("docs/M365_INTEGRATION.md")),
    ("Repository Organization", Path("docs/REPOSITORY_ORGANIZATION.md")),
    ("Repository Audit (2025-10-15)", Path("docs/REPOSITORY-AUDIT-2025-10-15.md")),
    ("Microsoft Support: Billing Dispute", Path("exports/support/YQAT-E6IU-BG7-PGB/MICROSOFT-SUPPORT-BILLING-DISPUTE-YQAT-E6IU-BG7-PGB.md")),
]


def render_markdown(doc, markdown_text):
    """
    Simple markdown renderer for python-docx.
    Handles: headings, bullet lists, numbered lists, code blocks, empty lines.
    """
    lines = markdown_text.split("\n")
    in_code_block = False
    code_lines = []
    list_indent = 0

    for line in lines:
        # Code block toggle
        if line.strip().startswith("```"):
            if in_code_block:
                # End code block
                code_text = "\n".join(code_lines)
                doc.add_paragraph(code_text, style="Intense Quote")
                code_lines = []
            in_code_block = not in_code_block
            continue

        if in_code_block:
            code_lines.append(line)
            continue

        # Headings
        if line.startswith("#"):
            level = 0
            while line.startswith("#"):
                level += 1
                line = line[1:]
            text = line.strip()

            if level == 1:
                doc.add_paragraph(text, style="Title")
            elif level == 2:
                doc.add_paragraph(text, style="Subtitle")
            elif level <= 9:
                doc.add_paragraph(text, style=f"Heading {level}")
            else:
                doc.add_paragraph(text, style="Heading 9")
            continue

        # Bullet lists
        if line.strip().startswith("- ") or line.strip().startswith("* "):
            text = line.strip()[2:]
            doc.add_paragraph(text, style="List Bullet")
            continue

        # Numbered lists (simple detection)
        if line.strip() and line.strip()[0].isdigit() and ". " in line:
            text = line.strip().split(". ", 1)[1] if ". " in line else line.strip()
            doc.add_paragraph(text, style="List Number")
            continue

        # Empty lines (skip)
        if not line.strip():
            continue

        # Normal paragraph
        doc.add_paragraph(line.strip())


def export_compendium(output_path, files):
    """Export multiple markdown files to a single DOCX compendium."""
    doc = Document()

    # Title page
    doc.add_paragraph("TEC-TGCR Compendium", style="Title")
    doc.add_paragraph("The Elidoras Codex — Theory of General Contextual Resonance", style="Subtitle")
    doc.add_paragraph("Agent Stack Documentation & Operations Guide")
    doc.add_page_break()

    # Process each file
    for title, path in files:
        print(f"[i] Adding: {title} ({path})")

        if not path.exists():
            print(f"[!] File not found, skipping: {path}")
            doc.add_paragraph(f"{title} — FILE NOT FOUND", style="Heading 1")
            doc.add_paragraph(f"Expected path: {path}")
            doc.add_page_break()
            continue

        # Add section heading
        doc.add_paragraph(title, style="Heading 1")

        # Read and render markdown
        markdown_content = path.read_text(encoding="utf-8")
        render_markdown(doc, markdown_content)

        # Page break after each document
        doc.add_page_break()

    # Save
    doc.save(output_path)
    print(f"\n[✓] Compendium exported to {output_path}")


def main():
    parser = argparse.ArgumentParser(description="Export TEC-TGCR documentation compendium to DOCX")
    parser.add_argument(
        "--output",
        "-o",
        type=Path,
        default=Path(__file__).parent.parent / "docs" / "exports" / "TEC_TGCR_COMPENDIUM.docx",
        help="Output path for DOCX file (default: docs/exports/TEC_TGCR_COMPENDIUM.docx)"
    )

    args = parser.parse_args()

    # Ensure output directory exists
    args.output.parent.mkdir(parents=True, exist_ok=True)

    # Export
    export_compendium(args.output, MANIFEST)


if __name__ == "__main__":
    main()
